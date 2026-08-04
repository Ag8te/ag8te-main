from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports"
OUT.mkdir(exist_ok=True)

DOCX_PATH = OUT / "AG8TE_E-Hailing_Checklist_Gap_Report.docx"
LOGO_PATH = ROOT / "static" / "assets" / "logo.jpeg"

NAVY = "17233A"
BLUE = "2563EB"
LIGHT_BLUE = "EAF2FF"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "667085"
DARK = "1F2937"
RED = "B42318"
RED_FILL = "FEECEB"
AMBER = "9A6700"
AMBER_FILL = "FFF4CC"
GREEN = "067647"
GREEN_FILL = "E8F7EF"
WHITE = "FFFFFF"


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_run(paragraph, text, *, bold=False, color=DARK, size=9.5):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_bullet(doc, text, color=DARK, after=3):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    add_run(p, text, color=color, size=9.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_status_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1500, 4980, 2880]
    headers = ["Priority", "Gap or requirement", "Recommended action"]
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_fill(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_run(p, text, bold=True, color=WHITE, size=9)

    for priority, gap, action in rows:
        cells = table.add_row().cells
        status_color, status_fill = {
            "Critical": (RED, RED_FILL),
            "High": (AMBER, AMBER_FILL),
            "Partial": (AMBER, AMBER_FILL),
            "Present": (GREEN, GREEN_FILL),
        }[priority]
        values = [priority, gap, action]
        for idx, value in enumerate(values):
            cell = cells[idx]
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if idx == 0:
                set_cell_fill(cell, status_fill)
            elif len(table.rows) % 2 == 0:
                set_cell_fill(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            add_run(p, value, bold=(idx == 0), color=status_color if idx == 0 else DARK, size=8.7)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.78)
section.right_margin = Inches(0.78)
section.header_distance = Inches(0.32)
section.footer_distance = Inches(0.32)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(9.5)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.1

for level, size, before, after in [(1, 15, 12, 6), (2, 11.5, 9, 4)]:
    style = styles[f"Heading {level}"]
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else BLUE)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

list_style = styles["List Bullet"]
list_style.font.name = "Arial"
list_style.font.size = Pt(9.5)
list_style.paragraph_format.left_indent = Inches(0.25)
list_style.paragraph_format.first_line_indent = Inches(-0.16)

# Running header
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(hp, "AG8TE  |  E-HAILING COMPLIANCE GAP REPORT", bold=True, color=MID_GRAY, size=7.5)

# Footer with dynamic page number
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(fp, "Confidential | Prepared for inspection readiness | Page ", color=MID_GRAY, size=7.5)
fld_char1 = OxmlElement("w:fldChar")
fld_char1.set(qn("w:fldCharType"), "begin")
instr_text = OxmlElement("w:instrText")
instr_text.set(qn("xml:space"), "preserve")
instr_text.text = "PAGE"
fld_char2 = OxmlElement("w:fldChar")
fld_char2.set(qn("w:fldCharType"), "end")
fp._p.append(fld_char1)
fp._p.append(instr_text)
fp._p.append(fld_char2)

# Cover/title block
if LOGO_PATH.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.05))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
add_run(p, "MZansiServe E-Hailing Platform", bold=True, color=NAVY, size=22)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
add_run(p, "Checklist Gap Assessment and Inspection Readiness Report", bold=True, color=BLUE, size=12)

meta = doc.add_table(rows=1, cols=3)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.autofit = False
meta_values = [
    ("Assessment basis", "Provided E-Hailing Platform Checklist"),
    ("Assessment date", date.today().strftime("%d %B %Y")),
    ("Status", "Action required before inspection"),
]
for idx, (label, value) in enumerate(meta_values):
    cell = meta.rows[0].cells[idx]
    set_cell_width(cell, 3120)
    set_cell_fill(cell, LIGHT_BLUE)
    set_cell_margins(cell, top=130, bottom=130)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_run(p, label.upper(), bold=True, color=MID_GRAY, size=7)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    add_run(p, value, bold=True, color=NAVY, size=8.5)

add_heading(doc, "Executive Summary", 1)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(7)
add_run(
    p,
    "AG8TE already contains several important e-hailing capabilities, including ride booking, "
    "scheduled trips, live driver location sharing, in-app messaging, ratings, payments, payouts, "
    "driver document collection, and administrative user management. However, the assessment identified "
    "several software, operational, and regulatory gaps that should be closed before a formal inspection "
    "or approval submission.",
    size=9.5,
)

callout = doc.add_table(rows=1, cols=1)
callout.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = callout.cell(0, 0)
set_cell_fill(cell, RED_FILL)
set_cell_margins(cell, top=140, bottom=140, start=170, end=170)
p = cell.paragraphs[0]
p.paragraph_format.space_after = Pt(2)
add_run(p, "IMMEDIATE PRIORITY", bold=True, color=RED, size=8)
p = cell.add_paragraph()
p.paragraph_format.space_after = Pt(0)
add_run(
    p,
    "Complete native push notifications, safety and SOS evidence, driver and vehicle compliance controls, "
    "navigation and heat-map tools, and required regulator reporting workflows.",
    bold=True,
    color=DARK,
    size=9.5,
)

add_heading(doc, "1. Critical Software and Operational Gaps", 1)
add_status_table(
    doc,
    [
        ("Critical", "Native mobile push notifications for ride requests, driver arrival, and trip progress.", "Implement and test push notifications on Android, iOS, and Huawei builds."),
        ("Critical", "Visible and fully demonstrated SOS/panic process, including weekly testing records.", "Connect the user-facing SOS control to the security provider and maintain a weekly test register."),
        ("Critical", "Vehicle emergency button linked to an installed tracking system.", "Select an approved tracker/provider and document installation and escalation procedures."),
        ("High", "Driver high-demand-area heat maps.", "Add demand aggregation and a driver-facing heat-map view."),
        ("High", "Driver navigation and route optimization.", "Integrate turn-by-turn navigation or approved external navigation handoff."),
        ("High", "Ability to accept a future ride while completing the current trip.", "Add controlled queued-ride acceptance with availability rules."),
        ("High", "Formal daily, weekly, and monthly driver earnings/performance reports.", "Create period-based reports with downloadable statements."),
        ("High", "Operational 24/7 help centre.", "Establish staffed support channels, escalation procedures, and response records."),
    ],
)

add_heading(doc, "2. Driver, Vehicle and Payment Gaps", 1)
add_status_table(
    doc,
    [
        ("Critical", "Driver police-clearance certificate and signed no-pending-criminal-case declaration.", "Add document upload, validation, expiry/review status, and administrator approval."),
        ("Critical", "Vehicle roadworthy certificate collection and validation.", "Add roadworthy evidence to vehicle onboarding and prevent expired vehicles from operating."),
        ("High", "Required platform name, address, and contact markings on vehicles.", "Define the approved vehicle-marking standard and retain photographic evidence."),
        ("Partial", "Recent driver photo identification.", "Enforce photo recency and administrator verification."),
        ("Partial", "Multiple payment methods are present, but cash and clearly defined in-app payment support require confirmation.", "Confirm approved payment methods, implement any missing method, and document settlement and refund flows."),
    ],
)

add_heading(doc, "3. Features Present but Requiring Evidence", 1)
for text in [
    "Realtime driver tracking exists, but it must be demonstrated on physical devices before inspection.",
    "Trip lifecycle and pickup confirmation exist, but the complete passenger and driver verification process should be documented.",
    "In-app notifications exist, but they do not replace native mobile push notifications.",
    "SOS backend integration exists, but the complete user-facing flow and emergency-provider response must be proven.",
    "Driver suspension exists, but the required remediation and regulator-notification workflow is incomplete.",
    "Yoco/PayPal payments and driver payouts exist, but evidence of security, reconciliation, refunds, and all required payment methods is needed.",
]:
    add_bullet(doc, text)

add_heading(doc, "4. Required Legal and Compliance Evidence", 1)
add_status_table(
    doc,
    [
        ("Critical", "Completed Form 9A and proof of payment.", "Prepare the signed application pack and bank receipt."),
        ("Critical", "Certified identity/company registration, proxy, address, and related supporting documents.", "Create a certified-document submission folder with an index."),
        ("Critical", "Tax clearance certificate and SARS PIN.", "Obtain current tax-compliance evidence."),
        ("Critical", "ICASA Type Approval certificates and RICA compliance evidence for all relevant equipment.", "Confirm every deployed device and retain certificates."),
        ("Critical", "Platform/operator agreement containing the checklist's required clauses.", "Have the agreement reviewed and signed by an authorised transport-compliance professional."),
        ("Critical", "Operating licence endorsement and approval for each platform.", "Document approval and endorsement for Google Play, Apple App Store, Huawei AppGallery, and any other platform."),
        ("High", "NPTR reporting process for suspensions and terminations, including the 24-hour requirement.", "Create an auditable notification workflow and retain submission evidence."),
        ("High", "Fourteen-day remediation process before cancellation.", "Implement case tracking, notices, deadlines, evidence review, and final decisions."),
        ("High", "Multi-platform registration and operator-notification process.", "Track each platform approval and record operator notifications."),
        ("High", "Cybersecurity, encryption, data protection, backup, and access-control evidence.", "Prepare a technical security pack with policies, architecture, tests, and incident procedures."),
        ("High", "Evidence that fare calculation and ticketing meet applicable transport requirements.", "Obtain a compliance review of fare calculation and interoperability requirements."),
    ],
)

add_heading(doc, "5. Recommended Action Plan", 1)
steps = [
    ("Phase 1 — Inspection blockers", "Complete push notifications, SOS and tracker evidence, police-clearance and roadworthy controls, Form 9A, tax evidence, and platform/operator agreements."),
    ("Phase 2 — Driver operations", "Implement navigation, demand heat maps, queued rides, formal earnings reports, vehicle markings, and 24/7 support procedures."),
    ("Phase 3 — Regulatory workflows", "Implement 14-day remediation, NPTR reporting, multi-platform registration tracking, operator-licence endorsement records, and licence-return procedures."),
    ("Phase 4 — Verification pack", "Run physical-device tests, security testing, payment reconciliation tests, weekly SOS tests, and compile screenshots, certificates, policies, and signed records."),
]
for idx, (label, description) in enumerate(steps, start=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_after = Pt(2)
    add_run(p, f"{idx}. {label}", bold=True, color=BLUE, size=10)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, description, size=9.5)

add_heading(doc, "Assessment Notes", 1)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
add_run(
    p,
    "This report compares the supplied checklist with functionality and evidence observed in the current "
    "AG8TE project. Operational activities and documents that are maintained outside the source-code "
    "repository must be confirmed separately. This report supports inspection preparation and does not "
    "replace advice from the relevant regulator, legal adviser, or transport-compliance professional.",
    color=MID_GRAY,
    size=8.5,
)

doc.core_properties.title = "AG8TE E-Hailing Checklist Gap Assessment"
doc.core_properties.subject = "Inspection readiness and compliance gap report"
doc.core_properties.author = "AG8TE"
doc.save(DOCX_PATH)
print(DOCX_PATH)
